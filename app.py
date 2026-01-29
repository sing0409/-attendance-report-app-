import streamlit as st
import pandas as pd
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt
import io

# 設定頁面配置
st.set_page_config(
    page_title="出勤報表自動彙整系統",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# 自訂 CSS 樣式
st.markdown("""
    <style>
    .main {
        padding: 2rem;
    }
    .header-title {
        font-size: 2.5rem;
        font-weight: 700;
        color: #1e40af;
        margin-bottom: 0.5rem;
    }
    .header-subtitle {
        font-size: 1rem;
        color: #64748b;
        margin-bottom: 2rem;
    }
    .upload-section {
        background: linear-gradient(135deg, #f0f9ff 0%, #f5f3ff 100%);
        padding: 2rem;
        border-radius: 12px;
        border: 2px dashed #3b82f6;
        margin-bottom: 2rem;
    }
    .success-message {
        background: #dcfce7;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #22c55e;
        margin: 1rem 0;
    }
    .error-message {
        background: #fee2e2;
        padding: 1rem;
        border-radius: 8px;
        border-left: 4px solid #ef4444;
        margin: 1rem 0;
    }
    </style>
""", unsafe_allow_html=True)

def format_date(date_str):
    """從日期字符串提取月/日"""
    match = re.search(r'(\d{4})-(\d{2})-(\d{2})', str(date_str))
    if match:
        return f"{int(match.group(2))}/{int(match.group(3))}"
    return str(date_str)

def process_attendance(df):
    """處理出勤紀錄並彙整數據"""
    summary = []
    
    for (emp_id, name, dept, title), group in df.groupby(['員工編號', '姓名', '部門', '職稱']):
        leaves_data = {}
        overtimes = []
        lates = []
        
        for _, row in group.iterrows():
            date_display = format_date(row['日期'])
            status = str(row['出勤狀況']) if pd.notna(row['出勤狀況']) else ""
            note = str(row['說明']) if pd.notna(row['說明']) else ""
            out_reg = str(row['外出登記']) if pd.notna(row['外出登記']) else ""
            late_min = row['遲到分鐘數']
            
            # 處理遲到
            late_val = 0
            if pd.notna(late_min):
                if isinstance(late_min, str):
                    m = re.search(r'(\d+)', late_min)
                    if m: late_val = int(m.group(1))
                else:
                    late_val = int(late_min)
            
            if late_val > 0:
                lates.append(f"{date_display} {late_val}M")
            
            # 處理請假與加班
            items = re.split(r'[\n\*]', status)
            for item in items:
                item = item.strip()
                if not item: continue
                
                # 處理加班
                if "[加班]" in item:
                    time_match = re.search(r'(\d{2}:\d{2})-(\d{2}:\d{2})', item)
                    if time_match:
                        start_str, end_str = time_match.groups()
                        start_dt = datetime.strptime(start_str, "%H:%M")
                        end_dt = datetime.strptime(end_str, "%H:%M")
                        hours = (end_dt - start_dt).seconds / 3600.0
                        
                        reason = ""
                        if out_reg:
                            dest_match = re.search(r'\[目的地\](.*?)(?:\n|\[|$)', out_reg, re.DOTALL)
                            desc_match = re.search(r'\[說明\](.*?)(?:\n|\[|$)', out_reg, re.DOTALL)
                            parts = []
                            if dest_match and dest_match.group(1).strip():
                                parts.append(dest_match.group(1).strip())
                            if desc_match and desc_match.group(1).strip():
                                parts.append(desc_match.group(1).strip())
                            if parts:
                                reason = " ".join(parts).replace('\n', ' ')
                        
                        if not reason and note:
                            clean_note = re.sub(r'\[\d{2}:\d{2}:\d{2}\]', '', note).strip()
                            if clean_note:
                                reason = clean_note.replace('\n', ' ')
                        
                        overtimes.append(f"{date_display} {hours:g}H {reason}".strip())
                
                # 處理請假
                leave_types = ["特休", "補休/調休", "生理假", "病假", "事假", "陪產檢及陪產假", "年休", "喪假", "公假"]
                for lt in leave_types:
                    if lt in item:
                        time_match = re.search(r'(\d{2}:\d{2})-(\d{2}:\d{2})', item)
                        if time_match:
                            start_str, end_str = time_match.groups()
                            start_dt = datetime.strptime(start_str, "%H:%M")
                            end_dt = datetime.strptime(end_str, "%H:%M")
                            hours = (end_dt - start_dt).seconds / 3600.0
                            
                            # 全天休假改為 8 小時
                            if hours >= 8.5:
                                hours = 8.0
                            
                            display_lt = lt
                            if lt == "補休/調休": display_lt = "補休"
                            if lt == "特休": display_lt = "年休"
                            
                            if display_lt not in leaves_data: leaves_data[display_lt] = []
                            leaves_data[display_lt].append(f"{date_display} {hours:g}H")
                        break

        leave_lines = []
        for k, v in leaves_data.items():
            leave_lines.append(f"{k}：{', '.join(v)}")
        leave_str = "\n".join(leave_lines)
        
        overtime_str = "\n".join(overtimes) if overtimes else ""
        late_str = "\n".join(lates) if lates else "無遲到"
        
        summary.append({
            "姓名": name,
            "職稱": title,
            "請假日期與時數": leave_str,
            "加班日期時數與原因": overtime_str,
            "遲到日期和遲到分鐘數": late_str
        })
    
    return pd.DataFrame(summary)

def create_word_report(summary_df, template_file):
    """將彙整數據填入 Word 範本"""
    doc = Document(template_file)
    
    if not doc.tables:
        return None
    
    table = doc.tables[0]
    
    # 找到欄位索引
    header_cells = [cell.text.strip() for cell in table.rows[0].cells]
    col_map = {}
    for i, text in enumerate(header_cells):
        if "姓名" in text: col_map["姓名"] = i
        elif "職稱" in text: col_map["職稱"] = i
        elif "假別" in text: col_map["請假日期與時數"] = i
        elif "加班" in text: col_map["加班日期時數與原因"] = i
        elif "遲到" in text: col_map["遲到日期和遲到分鐘數"] = i

    # 清空現有資料列
    for r in range(1, len(table.rows)):
        for cell in table.rows[r].cells:
            cell.text = ""

    # 填入資料
    for i, row_data in summary_df.iterrows():
        if i + 1 < len(table.rows):
            row = table.rows[i + 1]
        else:
            row = table.add_row()
            
        for col_name, col_idx in col_map.items():
            cell = row.cells[col_idx]
            cell.text = str(row_data.get(col_name, ""))
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    #run.font.name = '微軟正黑體'

    # 將檔案保存到記憶體
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# 主頁面
st.markdown('<div class="header-title">📊 出勤報表自動彙整系統</div>', unsafe_allow_html=True)
st.markdown('<div class="header-subtitle">社團法人中華民國更生少年關懷協會</div>', unsafe_allow_html=True)

st.markdown("""
本系統可自動彙整員工出勤紀錄，包含請假、加班及遲到資訊。
只需上傳 Excel 出勤紀錄與 Word 範本，即可快速生成完整的月報表。
""")

# 上傳區域
st.markdown('<div class="upload-section">', unsafe_allow_html=True)
col1, col2 = st.columns(2)

with col1:
    st.subheader("📁 上傳出勤紀錄")
    excel_file = st.file_uploader("選擇 Excel 檔案", type=["xlsx", "xls"], key="excel")

with col2:
    st.subheader("📄 上傳 Word 範本")
    word_file = st.file_uploader("選擇 Word 檔案", type=["docx"], key="word")

st.markdown('</div>', unsafe_allow_html=True)

# 處理邏輯
if excel_file and word_file:
    try:
        # 讀取 Excel
        df = pd.read_excel(excel_file, sheet_name='出勤紀錄')
        
        # 處理數據
        with st.spinner('正在彙整出勤資料...'):
            summary_df = process_attendance(df)
            output_file = create_word_report(summary_df, word_file)
        
        if output_file:
            st.markdown('<div class="success-message">✅ 報表彙整成功！</div>', unsafe_allow_html=True)
            
            # 顯示彙整摘要
            st.subheader("📋 彙整摘要")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("總員工數", len(summary_df))
            with col2:
                has_leave = (summary_df['請假日期與時數'] != '').sum()
                st.metric("有請假紀錄", has_leave)
            with col3:
                has_overtime = (summary_df['加班日期時數與原因'] != '').sum()
                st.metric("有加班紀錄", has_overtime)
            
            # 顯示詳細資料
            st.subheader("👥 員工出勤詳情")
            st.dataframe(summary_df, use_container_width=True)
            
            # 下載按鈕
            st.download_button(
                label="📥 下載彙整報表",
                data=output_file,
                file_name=f"出勤月報表_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        else:
            st.markdown('<div class="error-message">❌ Word 範本處理失敗，請確認檔案格式。</div>', unsafe_allow_html=True)
            
    except Exception as e:
        st.markdown(f'<div class="error-message">❌ 處理出錯：{str(e)}</div>', unsafe_allow_html=True)
        st.error(f"詳細錯誤：{str(e)}")
else:
    st.info("👆 請上傳 Excel 出勤紀錄和 Word 範本，系統會自動為您彙整報表。")
